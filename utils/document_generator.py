from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment
import os
from datetime import datetime
import re
import json

class DocumentGenerator:
    def __init__(self, template_path, output_dir):
        self.template_path = template_path
        self.output_dir = output_dir
        
    def generate_pass_request(self, data, request_type):
        """Генерация заявки на пропуск"""
        try:
            doc = Document(self.template_path)
            
            # Обработка данных
            processed_data = self._process_pass_request_data(data, request_type)
            
            # Замена тегов
            self._replace_tags_in_document(doc, processed_data)
            
            # Обработка таблиц с сотрудниками и транспортом
            self._process_tables(doc, data)
            
            filename = f"pass_request_{request_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(self.output_dir, filename)
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            print(f"Ошибка генерации документа: {e}")
            return None
    
    def _process_pass_request_data(self, data, request_type):
        """Обработка данных для заявки на пропуск"""
        processed = data.copy()
        
        # Обработка периода проезда
        if 'start_date' in data and 'end_date' in data:
            start = datetime.strptime(data['start_date'], '%Y-%m-%d').strftime('%d.%m.%Y')
            end = datetime.strptime(data['end_date'], '%Y-%m-%d').strftime('%d.%m.%Y')
            processed['ПЕРИОД ПРОЕЗДА'] = f"с {start} по {end}"
        
        # Обработка постов
        if 'posts' in data:
            posts = json.loads(data['posts'])
            processed['ПЕРЕЧЕНЬ ПОСТОВ'] = ', '.join(posts)
        
        return processed
    
    def _process_tables(self, doc, data):
        """Обработка таблиц с сотрудниками и транспортом"""
        # Здесь реализуется логика добавления/удаления строк в таблицах
        # в зависимости от выбранных сотрудников и транспорта
        pass
    
    def _replace_tags_in_document(self, doc, data):
        """Замена тегов во всем документе"""
        for paragraph in doc.paragraphs:
            self._replace_tags_in_paragraph(paragraph, data)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_tags_in_paragraph(paragraph, data)
    
    def _replace_tags_in_paragraph(self, paragraph, data):
        """Замена тегов в параграфе"""
        for key, value in data.items():
            tag = f"{{{{{key}}}}}"
            if tag in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(tag, str(value))
    
    def generate_ttn(self, data):
        """Генерация ТТН"""
        try:
            doc = Document(self.template_path)
            
            # Обработка данных ТТН
            processed_data = {
                'ДАТАСОСТАВЛЕНИЯ': datetime.strptime(data['date'], '%Y-%m-%d').strftime('%d.%m.%Y'),
                'НОМЕРТТН': data['number'],
                'ОРГАНИЗАЦИЯ ГРУЗООТПРАВИТЕЛЬ': data.get('sender_organization', ''),
                'ОРГАНИЗАЦИЯГРУЗОПОЛУЧАТЕЛЬ': data.get('receiver_organization', ''),
                'МЕСТА': data.get('places_count', ''),
                'МАССАГРУЗА': data.get('cargo_weight', ''),
                'АДРЕСПОГРУЗКИ': data.get('loading_address', ''),
                'АДРЕСРАЗГРУЗКИ': data.get('unloading_address', ''),
                'ВРЕМЯПОГРУЗКИ': data.get('loading_time', ''),
                'ГРУЗОТПРАВИТЕЛЬФИЗ': data.get('sender_individual', ''),
                'ПЕРЕВОЗЧИКФИЗ': data.get('carrier_individual', ''),
                'ГРУЗООТПРАВИТЕЛЬЮР': data.get('sender_legal', ''),
                'ТС': data.get('vehicle_info', ''),
                'НОМЕРПЛ': data.get('waybill_number', ''),
                'ПРИЦЕП': data.get('trailer_info', '')
            }
            
            # Обработка перечня груза
            if 'cargo_list' in data:
                cargo_items = json.loads(data['cargo_list'])
                processed_data['НАИМЕНОВАНИЕГРУЗА'] = ', '.join([item['name'] for item in cargo_items])
            
            self._replace_tags_in_document(doc, processed_data)
            
            filename = f"ttn_{data['number']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(self.output_dir, filename)
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            print(f"Ошибка генерации ТТН: {e}")
            return None
    
    def generate_shift_request(self, data, request_type):
        """Генерация заявки на перевахтовку"""
        try:
            wb = load_workbook(self.template_path)
            ws = wb.active
            
            # Обработка данных
            processed_data = self._process_shift_request_data(data, request_type)
            
            # Замена тегов в ячейках
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        for key, value in processed_data.items():
                            tag = f"{{{{{key}}}}}"
                            if tag in str(cell.value):
                                cell.value = str(cell.value).replace(tag, str(value))
            
            # Обработка таблицы с сотрудниками
            self._process_employee_table(ws, data)
            
            filename = f"shift_request_{request_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            output_path = os.path.join(self.output_dir, filename)
            wb.save(output_path)
            return output_path
            
        except Exception as e:
            print(f"Ошибка генерации Excel документа: {e}")
            return None
    
    def _process_shift_request_data(self, data, request_type):
        """Обработка данных для заявки на перевахтовку"""
        processed = {}
        
        if request_type == 'charter':
            processed.update({
                'ОТКУДА': data.get('departure_airport', ''),
                'КУДА': data.get('arrival_airport', ''),
                'ДАТА ВЫЛЕТА': datetime.strptime(data['flight_date'], '%Y-%m-%d').strftime('%d.%m.%Y') if data.get('flight_date') else '',
                'ОТКУДА АВТО': data.get('auto_delivery_from', ''),
                'КУДА АВТО': data.get('auto_delivery_to', '')
            })
        elif request_type == 'regular':
            processed.update({
                'ВЫЛЕТ': data.get('departure_airport', ''),
                'ПРИЛЕТ': data.get('arrival_airport', ''),
                'ДАТА ВЫЛЕТА': datetime.strptime(data['flight_date'], '%Y-%m-%d').strftime('%d.%m.%Y') if data.get('flight_date') else '',
                'СТОИМОСТЬ': data.get('preliminary_cost', ''),
                'НОМЕР РЕЙСА': data.get('flight_number', '')
            })
        
        return processed
    
    def _process_employee_table(self, ws, data):
        """Обработка таблицы с сотрудниками в Excel"""
        # Реализация добавления строк с сотрудниками в таблицу
        pass

class ElectricityTracker:
    """Класс для учета электроэнергии"""
    
    @staticmethod
    def update_readings(file_path, date, previous_bpo, previous_dormitory, current_bpo, current_dormitory):
        """Обновление показаний электроэнергии в Excel файле"""
        try:
            wb = load_workbook(file_path)
            
            # Создаем новый лист с названием месяца и года
            month_name = date.strftime('%B %Y')
            new_sheet = wb.copy_worksheet(wb.worksheets[-1])
            new_sheet.title = month_name
            
            # Обновляем показания
            new_sheet['F5'] = previous_bpo
            new_sheet['G5'] = current_bpo
            new_sheet['F6'] = previous_dormitory
            new_sheet['G6'] = current_dormitory
            
            wb.save(file_path)
            return True
            
        except Exception as e:
            print(f"Ошибка обновления показаний: {e}")
            return False